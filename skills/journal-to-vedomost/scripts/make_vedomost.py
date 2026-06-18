"""
Шаблон: Генерация зачётной ведомости (.docx) из данных журнала.
Заполните переменные META и STUDENTS, затем запустите скрипт.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ═══════════════════════════════════════════════════════════════
# ЗАПОЛНИТЕ ЭТИ ДАННЫЕ
# ═══════════════════════════════════════════════════════════════

OUT_PATH = r"Зачётные ведомости\ГРУППА_ДИСЦИПЛИНА_зачётная_ведомость.docx"

META = {
    "discipline":    "УП.XX - Учебная практика",   # из заголовка журнала
    "course":        "1",                            # курс
    "group":         "ХХИСд-ХХ",                    # группа
    "semester":      "1",                            # 1 = сен-янв, 2 = фев-июн
    "spec":          "09.02.07 - Информационные системы и программирование",
    "date":          "ДД.ММ.ГГГГ г.",               # последняя дата в журнале
    "teacher":       "Фамилия Имя Отчество",         # ФИО преподавателя
    "teacher_short": "Фамилия И.О.",
}

# Список студентов: (номер, "ФИО", "оценка")
# Оценка: "5", "4", "3", "2" или "А/З"
STUDENTS = [
    (1,  "Фамилия Имя Отчество",   "5"),
    (2,  "Фамилия Имя Отчество",   "А/З"),
    # ... добавьте остальных
]

# ═══════════════════════════════════════════════════════════════
# КОД ГЕНЕРАЦИИ (не изменять)
# ═══════════════════════════════════════════════════════════════

def set_cell(cell, text, bold=False, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             v_align=WD_ALIGN_VERTICAL.CENTER,
             bg_hex=None):
    cell.vertical_alignment = v_align
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size)
    if bg_hex:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg_hex)
        tcPr.append(shd)

def add_para(doc, text, bold=False, size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    return p

doc = Document()
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(1.5)

add_para(doc, "ЗАЧЁТНАЯ ВЕДОМОСТЬ", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para(doc, f"по учебному предмету/дисциплине/МДК: {META['discipline']}",
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_para(doc, f"Курс: {META['course']}        Группа: {META['group']}        Семестр: {META['semester']}", size=11)
add_para(doc, f"Код и наименование специальности: {META['spec']}", size=11)
add_para(doc, f"Дата проведения: {META['date']}", size=11)
add_para(doc, f"ФИО преподавателя: {META['teacher']}", size=11, space_after=8)

col_widths = [Cm(1.5), Cm(9), Cm(2.5), Cm(5)]
table = doc.add_table(rows=1 + len(STUDENTS), cols=4)
table.style     = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, title in enumerate(["№ п/п", "ФИО обучающегося", "Оценка", "Подпись преподавателя"]):
    set_cell(table.rows[0].cells[i], title, bold=True, size=10, bg_hex="BDD7EE")

for num, fio, grade in STUDENTS:
    row = table.rows[num]
    az  = grade == "А/З"
    bg  = "FFE0E0" if az else None
    set_cell(row.cells[0], str(num), size=10, bg_hex=bg)
    set_cell(row.cells[1], fio, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, bg_hex=bg)
    set_cell(row.cells[2], grade, bold=az, size=10, bg_hex=bg)
    set_cell(row.cells[3], "", size=10, bg_hex=bg)

for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

doc.add_paragraph()
add_para(doc, f"Преподаватель: ____________________  {META['teacher_short']}", size=11)

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Готово: {OUT_PATH}")
